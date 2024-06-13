from docx import Document

# Create a new Document
doc = Document()

# Add a title
doc.add_heading('The Weather Forecast', 0)

# Add three lines of text
doc.add_paragraph('In Israel the temperatures are high.')
doc.add_paragraph('In France the temperatures are low.')
doc.add_paragraph('In Russia temperatures are below 0.')

# Save the document
doc.save('weather_document.docx')

# Load the document to verify the text
document = Document('weather_document.docx')

# Verify the text in the first paragraph
print(document.paragraphs[0].text)  # Output: This is the first line of text.
